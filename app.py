from flask import Flask, render_template, request, jsonify, send_file
from flask_cors import CORS
from deep_translator import GoogleTranslator
from langdetect import detect
import speech_recognition as sr
import base64, io, os, json
from datetime import datetime
from fpdf import FPDF
from docx import Document

app = Flask(__name__)
CORS(app)
DB_FILE = 'history.json'

# ── History helpers ──
def load_history():
    if os.path.exists(DB_FILE):
        with open(DB_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_history(entry):
    data = load_history()
    data.insert(0, entry)
    if len(data) > 100:
        data = data[:100]
    with open(DB_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

# ── Routes ──
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/translate', methods=['POST'])
def translate():
    try:
        data = request.get_json()
        text = data.get('text', '')
        target_lang = data.get('target_lang', 'hi')
        if not text:
            return jsonify({'success': False, 'text': 'No text to translate'})
        translated = GoogleTranslator(source='auto', target=target_lang).translate(text)
        return jsonify({'success': True, 'text': translated})
    except Exception as e:
        return jsonify({'success': False, 'text': 'Error: ' + str(e)})

@app.route('/detect_language', methods=['POST'])
def detect_language():
    try:
        data = request.get_json()
        text = data.get('text', '')
        if not text or len(text) < 5:
            return jsonify({'success': False, 'lang': 'unknown'})
        lang = detect(text)
        return jsonify({'success': True, 'lang': lang})
    except Exception as e:
        return jsonify({'success': False, 'lang': 'unknown'})

@app.route('/save_history', methods=['POST'])
def save_history_route():
    try:
        data = request.get_json()
        entry = {
            'id': datetime.now().strftime('%Y%m%d%H%M%S'),
            'text': data.get('text', ''),
            'translation': data.get('translation', ''),
            'lang': data.get('lang', 'English'),
            'duration': data.get('duration', 0),
            'word_count': data.get('word_count', 0),
            'char_count': data.get('char_count', 0),
            'timestamp': datetime.now().strftime('%d/%m/%Y %H:%M')
        }
        save_history(entry)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False})

@app.route('/get_history', methods=['GET'])
def get_history():
    return jsonify(load_history())

@app.route('/clear_history', methods=['POST'])
def clear_history():
    try:
        if os.path.exists(DB_FILE):
            os.remove(DB_FILE)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False})

# ── Export PDF ──
@app.route('/export_pdf', methods=['POST'])
def export_pdf():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        translation = data.get('translation', '').strip()

        if not text:
            return jsonify({'error': 'No text to export'}), 400

        pdf = FPDF()
        pdf.add_page()

        # Check if NotoSans font exists for unicode support
        font_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'NotoSans-Regular.ttf')
        has_unicode = os.path.exists(font_path)
        if has_unicode:
            pdf.add_font('NotoSans', '', font_path)
            
            # Add fallback fonts for Indian languages
            fallback_fonts = []
            for font_name in ['NotoSansDevanagari', 'NotoSansKannada', 'NotoSansTamil', 'NotoSansTelugu', 'NotoSansMalayalam', 'NotoSansBengali']:
                f_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), f'{font_name}-Regular.ttf')
                if os.path.exists(f_path):
                    try:
                        pdf.add_font(font_name, '', f_path)
                        fallback_fonts.append(font_name)
                    except:
                        pass
            if fallback_fonts and hasattr(pdf, 'set_fallback_fonts'):
                pdf.set_fallback_fonts(fallback_fonts)

        # Title
        pdf.set_font('Helvetica', 'B', 22)
        pdf.set_text_color(108, 99, 255)
        pdf.cell(0, 14, 'VoiceScribe', ln=True, align='C')

        pdf.set_font('Helvetica', '', 11)
        pdf.set_text_color(150, 150, 180)
        pdf.cell(0, 8, 'AI-Based Speech to Text System for Indian Languages', ln=True, align='C')
        pdf.cell(0, 8, 'Generated: ' + datetime.now().strftime('%d %B %Y, %H:%M'), ln=True, align='C')
        pdf.ln(4)

        # Divider
        pdf.set_draw_color(108, 99, 255)
        pdf.set_line_width(0.8)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(8)

        # Transcription section
        pdf.set_font('Helvetica', 'B', 14)
        pdf.set_text_color(30, 30, 80)
        pdf.cell(0, 10, 'Transcription', ln=True)
        pdf.ln(2)
        pdf.set_text_color(50, 50, 50)
        if has_unicode:
            pdf.set_font('NotoSans', '', 12)
            pdf.multi_cell(0, 8, text)
        else:
            pdf.set_font('Helvetica', '', 12)
            safe_text = text.encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 8, safe_text)

        # Translation section
        if translation:
            pdf.ln(6)
            pdf.set_draw_color(200, 200, 230)
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(6)
            pdf.set_font('Helvetica', 'B', 14)
            pdf.set_text_color(30, 30, 80)
            pdf.cell(0, 10, 'Translation', ln=True)
            pdf.ln(2)
            pdf.set_text_color(50, 50, 50)
            if has_unicode:
                pdf.set_font('NotoSans', '', 12)
                pdf.multi_cell(0, 8, translation)
            else:
                pdf.set_font('Helvetica', '', 12)
                safe_trans = translation.encode('latin-1', 'replace').decode('latin-1')
                pdf.multi_cell(0, 8, safe_trans)

        # Footer
        pdf.ln(10)
        pdf.set_font('Helvetica', 'I', 9)
        pdf.set_text_color(180, 180, 200)
        pdf.cell(0, 8, 'Generated by VoiceScribe - RNS Institute of Technology, Bengaluru', align='C')

        buf = io.BytesIO()
        buf.write(bytes(pdf.output()))
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf',
                        as_attachment=True,
                        download_name='voicescribe_transcription.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ── Export Word ──
@app.route('/export_word', methods=['POST'])
def export_word():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        translation = data.get('translation', '').strip()

        if not text:
            return jsonify({'error': 'No text to export'}), 400

        doc = Document()

        # Title
        title = doc.add_heading('VoiceScribe', 0)
        title.alignment = 1
        sub = doc.add_paragraph('AI-Based Speech to Text System for Indian Languages')
        sub.alignment = 1
        date_p = doc.add_paragraph('Generated: ' + datetime.now().strftime('%d %B %Y, %H:%M'))
        date_p.alignment = 1
        doc.add_paragraph('')

        # Transcription
        doc.add_heading('Transcription', level=1)
        doc.add_paragraph(text)

        # Translation
        if translation:
            doc.add_paragraph('')
            doc.add_heading('Translation', level=1)
            doc.add_paragraph(translation)

        # Footer
        doc.add_paragraph('')
        footer = doc.add_paragraph('Generated by VoiceScribe - RNS Institute of Technology, Bengaluru')
        footer.alignment = 1

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True,
                        download_name='voicescribe_transcription.docx')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ── Export TXT ──
@app.route('/export_txt', methods=['POST'])
def export_txt():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        translation = data.get('translation', '').strip()

        if not text:
            return jsonify({'error': 'No text to export'}), 400

        content = 'VOICESCRIBE - TRANSCRIPTION\n'
        content += '=' * 40 + '\n'
        content += 'Generated: ' + datetime.now().strftime('%d %B %Y, %H:%M') + '\n'
        content += '=' * 40 + '\n\n'
        content += 'TRANSCRIPTION:\n' + text + '\n'
        if translation:
            content += '\nTRANSLATION:\n' + translation + '\n'
        content += '\n' + '-' * 40 + '\n'
        content += 'Generated by VoiceScribe - RNSIT Bengaluru\n'

        buf = io.BytesIO(content.encode('utf-8'))
        buf.seek(0)
        return send_file(buf, mimetype='text/plain',
                        as_attachment=True,
                        download_name='voicescribe_transcription.txt')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ── Transcribe Audio File ──
@app.route('/transcribe_file', methods=['POST'])
def transcribe_file():
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'text': 'No data received'})

        audio_base64 = data.get('audio', '')
        lang = data.get('lang', 'en-IN')

        if not audio_base64:
            return jsonify({'success': False, 'text': 'No audio data received'})

        audio_bytes = base64.b64decode(audio_base64)
        recognizer = sr.Recognizer()

        # Try WAV format first
        try:
            audio_io = io.BytesIO(audio_bytes)
            with sr.AudioFile(audio_io) as source:
                recognizer.adjust_for_ambient_noise(source)
                audio = recognizer.record(source)
            text = recognizer.recognize_google(audio, language=lang)
            return jsonify({'success': True, 'text': text})
        except Exception:
            pass

        # Try pydub conversion for mp3/m4a/ogg
        try:
            from pydub import AudioSegment
            audio_io = io.BytesIO(audio_bytes)
            audio_segment = AudioSegment.from_file(audio_io)
            audio_segment = audio_segment.set_channels(1).set_frame_rate(16000)
            wav_io = io.BytesIO()
            audio_segment.export(wav_io, format='wav')
            wav_io.seek(0)
            with sr.AudioFile(wav_io) as source:
                audio = recognizer.record(source)
            text = recognizer.recognize_google(audio, language=lang)
            return jsonify({'success': True, 'text': text})
        except Exception as e2:
            return jsonify({
                'success': False,
                'text': 'Please upload a .WAV file. MP3/M4A needs ffmpeg. Error: ' + str(e2)
            })

    except sr.UnknownValueError:
        return jsonify({'success': False, 'text': 'Could not understand audio. Please speak clearly.'})
    except sr.RequestError:
        return jsonify({'success': False, 'text': 'Internet error. Check your connection.'})
    except Exception as e:
        return jsonify({'success': False, 'text': 'Error: ' + str(e)})

# ── Summarize Text ──
@app.route('/summarize', methods=['POST'])
def summarize_text():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        if not text or len(text.split()) < 10:
            return jsonify({'success': False, 'error': 'Text is too short to summarize'})
        
        from sumy.parsers.plaintext import PlaintextParser
        from sumy.nlp.tokenizers import Tokenizer
        from sumy.summarizers.lsa import LsaSummarizer
        import nltk
        nltk.download('punkt_tab', quiet=True)
        
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, 3)
        
        summary_text = "• " + "\n• ".join([str(sentence) for sentence in summary])
        return jsonify({'success': True, 'summary': summary_text})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# ── Smart Format ──
@app.route('/format_text', methods=['POST'])
def format_text():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        if not text:
            return jsonify({'success': False, 'error': 'No text provided'})
            
        text = text.capitalize()
        words = text.split()
        formatted = []
        for i, word in enumerate(words):
            if word.lower() in ['and', 'but', 'because', 'however', 'therefore'] and i > 0 and len(formatted) > 0 and not formatted[-1].endswith(','):
                formatted[-1] += ','
            formatted.append(word)
            
        res = ' '.join(formatted)
        if res and not res[-1] in ['.', '!', '?']:
            res += '.'
            
        return jsonify({'success': True, 'text': res})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# ── Export SRT ──
@app.route('/export_srt', methods=['POST'])
def export_srt():
    try:
        data = request.get_json()
        text = data.get('text', '').strip()
        duration = float(data.get('duration', 0))
        
        if not text:
            return jsonify({'error': 'No text to export'}), 400
            
        words = text.split()
        if duration == 0:
            duration = len(words) / 2.5
            
        chunk_size = 7
        chunks = [words[i:i + chunk_size] for i in range(0, len(words), chunk_size)]
        
        time_per_chunk = duration / len(chunks) if chunks else 0
        
        srt_content = ""
        current_time = 0.0
        
        def format_time(seconds):
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            secs = int(seconds % 60)
            millis = int((seconds % 1) * 1000)
            return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
            
        for i, chunk in enumerate(chunks):
            start_time = format_time(current_time)
            end_time = format_time(current_time + time_per_chunk)
            
            srt_content += f"{i+1}\n"
            srt_content += f"{start_time} --> {end_time}\n"
            srt_content += " ".join(chunk) + "\n\n"
            
            current_time += time_per_chunk
            
        buf = io.BytesIO(srt_content.encode('utf-8'))
        buf.seek(0)
        return send_file(buf, mimetype='text/plain',
                        as_attachment=True,
                        download_name='voicescribe_subtitles.srt')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)